from __future__ import annotations

"""
Document extraction for ClarityAI.

Adds:
- DOCX support via python-docx (graceful fallback if not installed).
- Markdown-aware light cleanup (preserves heading structure).
- HTML-to-text via stdlib (no extra deps) for .html / .htm uploads.
- Section-per-paragraph for long text files so chunker can split sensibly.
"""

import csv
import hashlib
import html
import io
import json
import logging
import mimetypes
import re
from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path

from pypdf import PdfReader

logger = logging.getLogger(__name__)


TEXT_LIKE_SUFFIXES = {".txt", ".md", ".markdown", ".rst", ".log", ".yaml", ".yml", ".ini", ".cfg"}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def sanitize_filename(name: str) -> str:
    cleaned = re.sub(r"[^a-zA-Z0-9._-]+", "-", name.strip())
    cleaned = cleaned.strip(".-")
    return cleaned or "document"


@dataclass(slots=True)
class ExtractedDocument:
    title: str
    source_name: str
    mime_type: str
    raw_text: str
    sections: list[dict]


def guess_mime_type(path: Path, explicit: str | None = None) -> str:
    if explicit:
        return explicit
    guessed, _ = mimetypes.guess_type(str(path))
    return guessed or "text/plain"


def compute_checksum(raw: bytes) -> str:
    return hashlib.sha256(raw).hexdigest()


def _clean_extracted_text(text: str) -> str:
    text = (text or "").replace("\u00a0", " ")
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


# ---------------------------------------------------------------------------
# Specific extractors
# ---------------------------------------------------------------------------


def _extract_pdf(path: Path, source_name: str, mime_type: str) -> ExtractedDocument:
    reader = PdfReader(str(path))
    sections: list[dict] = []
    all_parts: list[str] = []
    for page_number, page in enumerate(reader.pages, start=1):
        try:
            text = _clean_extracted_text(page.extract_text() or "")
        except Exception:
            text = ""
        if not text:
            continue
        sections.append({"text": text, "page_label": str(page_number)})
        all_parts.append(f"Page {page_number}\n{text}")
    raw_text = "\n\n".join(all_parts).strip()
    return ExtractedDocument(
        title=path.stem,
        source_name=source_name,
        mime_type=mime_type,
        raw_text=raw_text,
        sections=sections,
    )


def _extract_csv(path: Path, source_name: str, mime_type: str) -> ExtractedDocument:
    rows: list[str] = []
    with path.open("r", encoding="utf-8", errors="ignore", newline="") as handle:
        reader = csv.reader(handle)
        for row in reader:
            cells = [cell.strip() for cell in row if cell and cell.strip()]
            if cells:
                rows.append(" | ".join(cells))
    raw_text = _clean_extracted_text("\n".join(rows))
    return ExtractedDocument(
        title=path.stem,
        source_name=source_name,
        mime_type=mime_type,
        raw_text=raw_text,
        sections=[{"text": raw_text, "page_label": None}],
    )


def _extract_json(path: Path, source_name: str, mime_type: str) -> ExtractedDocument:
    with path.open("r", encoding="utf-8", errors="ignore") as handle:
        data = json.load(handle)
    raw_text = _clean_extracted_text(json.dumps(data, indent=2, ensure_ascii=False))
    return ExtractedDocument(
        title=path.stem,
        source_name=source_name,
        mime_type=mime_type,
        raw_text=raw_text,
        sections=[{"text": raw_text, "page_label": None}],
    )


def _extract_text_like(path: Path, source_name: str, mime_type: str) -> ExtractedDocument:
    raw_text = _clean_extracted_text(path.read_text(encoding="utf-8", errors="ignore"))
    # Split on blank lines so each section is a paragraph block. The chunker
    # will further break long sections.
    blocks = [b.strip() for b in re.split(r"\n{2,}", raw_text) if b.strip()]
    sections = [{"text": b, "page_label": None} for b in blocks] or [
        {"text": raw_text, "page_label": None}
    ]
    return ExtractedDocument(
        title=path.stem,
        source_name=source_name,
        mime_type=mime_type,
        raw_text=raw_text,
        sections=sections,
    )


def _extract_docx(path: Path, source_name: str, mime_type: str) -> ExtractedDocument:
    """Best-effort DOCX extractor.

    Uses python-docx if available; otherwise falls back to reading the
    document.xml inside the .docx zip and stripping tags.
    """
    try:
        from docx import Document  # type: ignore

        doc = Document(str(path))
        paragraphs: list[str] = []
        for p in doc.paragraphs:
            txt = (p.text or "").strip()
            if txt:
                paragraphs.append(txt)
        # Pull table text in too — lots of docs put data in tables.
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        raw_text = _clean_extracted_text("\n\n".join(paragraphs))
    except Exception:
        # Fallback: parse the underlying XML
        import zipfile

        try:
            with zipfile.ZipFile(str(path)) as zf:
                with zf.open("word/document.xml") as fp:
                    xml = fp.read().decode("utf-8", errors="ignore")
            text = re.sub(r"<w:p[^>]*>", "\n\n", xml)
            text = re.sub(r"<[^>]+>", "", text)
            raw_text = _clean_extracted_text(html.unescape(text))
        except Exception as exc:
            logger.warning("DOCX fallback parse failed for %s: %s", path, exc)
            raw_text = ""

    blocks = [b.strip() for b in re.split(r"\n{2,}", raw_text) if b.strip()] or [raw_text]
    sections = [{"text": b, "page_label": None} for b in blocks if b]
    return ExtractedDocument(
        title=path.stem,
        source_name=source_name,
        mime_type=mime_type,
        raw_text=raw_text,
        sections=sections,
    )


class _HTMLTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self._parts: list[str] = []
        self._skip_stack: list[str] = []

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag in {"script", "style", "noscript"}:
            self._skip_stack.append(tag)
        if tag in {"p", "br", "div", "li", "h1", "h2", "h3", "h4", "h5", "h6", "tr"}:
            self._parts.append("\n")
        if tag == "td":
            self._parts.append(" | ")

    def handle_endtag(self, tag: str) -> None:
        if self._skip_stack and self._skip_stack[-1] == tag:
            self._skip_stack.pop()
        if tag in {"p", "div", "li", "h1", "h2", "h3", "h4", "h5", "h6", "tr"}:
            self._parts.append("\n")

    def handle_data(self, data: str) -> None:
        if self._skip_stack:
            return
        if data.strip():
            self._parts.append(data)

    @property
    def text(self) -> str:
        return "".join(self._parts)


def _extract_html(path: Path, source_name: str, mime_type: str) -> ExtractedDocument:
    raw = path.read_text(encoding="utf-8", errors="ignore")
    parser = _HTMLTextExtractor()
    try:
        parser.feed(raw)
    except Exception:
        pass
    raw_text = _clean_extracted_text(parser.text)
    blocks = [b.strip() for b in re.split(r"\n{2,}", raw_text) if b.strip()] or [raw_text]
    sections = [{"text": b, "page_label": None} for b in blocks if b]
    return ExtractedDocument(
        title=path.stem,
        source_name=source_name,
        mime_type=mime_type,
        raw_text=raw_text,
        sections=sections,
    )


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------


def extract_document(path: Path, source_name: str, content_type: str | None = None) -> ExtractedDocument:
    mime_type = guess_mime_type(path, explicit=content_type)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(path, source_name=source_name, mime_type=mime_type)
    if suffix == ".csv":
        return _extract_csv(path, source_name=source_name, mime_type=mime_type)
    if suffix == ".json":
        return _extract_json(path, source_name=source_name, mime_type=mime_type)
    if suffix == ".docx":
        return _extract_docx(path, source_name=source_name, mime_type=mime_type)
    if suffix in {".html", ".htm"}:
        return _extract_html(path, source_name=source_name, mime_type=mime_type)
    if suffix in TEXT_LIKE_SUFFIXES or mime_type.startswith("text/"):
        return _extract_text_like(path, source_name=source_name, mime_type=mime_type)

    # Unknown binary — best-effort treat as text.
    return _extract_text_like(path, source_name=source_name, mime_type="text/plain")